"""Записывает введенные пользователем данные в docx формат"""

import docx
from docx import Document
from docx.shared import Pt
from services.services import read_json
import datetime


class Naryad:
    """Записывает введенные пользователем данные в docx формат"""
    doc = Document('data/original_naryad/original_naryad.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(10)
    coordinates_paste_tables = read_json("data/coordinates.json")["tables"]

    def __init__(self, type_naryad,
                 issuing,
                 head_works,
                 manufacturer_works,
                 brigade_members,
                 date_start,
                 time_start,
                 date_end,
                 time_end,
                 date_issued,
                 time_issued):

        self.type_naryad = type_naryad
        self.issuing = issuing
        self.head_works = head_works
        self.manufacturer_works = manufacturer_works
        self.brigade_members = brigade_members
        self.date_start = date_start
        self.time_start = time_start
        self.date_end = date_end
        self.time_end = time_end
        self.date_issued = date_issued
        self.time_issued = time_issued

    def paste_blank_one_value(self, coordinate: list, value: str) -> None:
        """Записывает текст по указанным координатам"""
        Naryad.doc.tables[coordinate[0]].cell(coordinate[1], coordinate[2]).text = value

    def paste_blank_name(self, coordinates: list, value: str) -> None:
        """Записывает данные по указанным координатам"""
        Naryad.doc.tables[coordinates[0][0]].cell(coordinates[0][1], coordinates[0][2]).text = value
        value = value.split(',')[0]
        coordinates = coordinates[1:]
        for coordinate in coordinates:
            Naryad.doc.tables[coordinate[0]].cell(coordinate[1], coordinate[2]).text = value

    def paste_blank_name_brigade_members(self, coordinates: list, value: list) -> None:
        """Записывает членов бригады по указанным координатам"""
        worker = ", ".join(value)
        Naryad.doc.tables[coordinates[0][0]].cell(coordinates[0][1], coordinates[0][2]).text = worker
        coordinates = coordinates[1:]
        for c, w in enumerate(value):
            for coordinate in coordinates:
                Naryad.doc.tables[coordinate[0]].cell(coordinate[1]+c, coordinate[2]).text = w.split(',')[0]

    def record_word(self) -> None:
        """Основная программа записи распределяющая значения в зависимости от польховательского ввода"""
        Naryad.paste_blank_name(self, Naryad.coordinates_paste_tables["issuing"], self.issuing)
        Naryad.paste_blank_one_value(self, Naryad.coordinates_paste_tables["type_naryad"], self.type_naryad)
        Naryad.paste_blank_one_value(self, Naryad.coordinates_paste_tables["date_start"], self.date_start)
        Naryad.paste_blank_one_value(self, Naryad.coordinates_paste_tables["time_start"], self.time_start)
        Naryad.paste_blank_one_value(self, Naryad.coordinates_paste_tables["date_end"], self.date_end)
        Naryad.paste_blank_one_value(self, Naryad.coordinates_paste_tables["time_end"], self.time_end)
        Naryad.paste_blank_one_value(self, Naryad.coordinates_paste_tables["date_issued"], self.date_issued)
        Naryad.paste_blank_one_value(self, Naryad.coordinates_paste_tables["time_issued"], self.time_issued)
        if self.head_works:
            Naryad.paste_blank_name(self, Naryad.coordinates_paste_tables["head_works"], self.head_works)
            Naryad.paste_blank_name(self, Naryad.coordinates_paste_tables["manufacturer_works_with_head_works"], self.manufacturer_works)
            Naryad.paste_blank_name_brigade_members(self, Naryad.coordinates_paste_tables["brigade_members_with_head_works"], self.brigade_members)
        else:
            Naryad.paste_blank_name(self, Naryad.coordinates_paste_tables["manufacturer_works_without_head_works"],
                               self.manufacturer_works)
            Naryad.paste_blank_name_brigade_members(self, Naryad.coordinates_paste_tables["brigade_members_without_head_works"], self.brigade_members)
        date_save_name = datetime.datetime.now().strftime('%d_%m_%y_%H-%M')
        Naryad.doc.save(f"data/fin_naryad/{date_save_name}.docx")





