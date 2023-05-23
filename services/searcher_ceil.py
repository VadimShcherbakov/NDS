# полноценный рабочий код
from docx import Document


# последовательность всех таблиц документа
def serch_need_ceil():
    doc = Document(r'data_services/serarch_position_naryad.docx')
    list_coordinates = []
    for num_table in range(len(doc.tables)):
        need_table = doc.tables[num_table]
        for num_row, row in enumerate(need_table.rows):
            for num_cell, cell in enumerate(row.cells):
                if cell.text == 'time':
                    print(f"таблица {num_table}, строка {num_row}, столбец {num_cell}")
                    list_coordinates.append([num_table, num_row, num_cell])
                    break
    return list_coordinates

print(serch_need_ceil())
